from django.shortcuts import render
from contact.models import Application, ApplicationCreate, Tuman, HujumTuri
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from openpyxl import Workbook
from openpyxl.writer.excel import save_virtual_workbook
from .views import login_decorator
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.styles import Border, Side
import os
from docx import Document
from django.http import FileResponse, HttpResponse
from django.conf import settings
import mimetypes


@login_decorator
def tablitsa(request):
    application = Application.objects.select_related('hujumturi', 'district', 'app_create').all()

    first_name = request.GET.get('first_name', '')
    last_name = request.GET.get('last_name', '')
    hujumturi_id = request.GET.get('hujum', '')
    district_id = request.GET.get('tuman', '')
    status = request.GET.get('status', '')

    if first_name:
        application = application.filter(first_name__icontains=first_name)
    if last_name:
        application = application.filter(last_name__icontains=last_name)
    if district_id:
        application = application.filter(district_id=district_id)
    if status:
        application = application.filter(app_create__status=status)  # Adjusted filter here
    if hujumturi_id:
        application = application.filter(hujumturi_id=hujumturi_id)

    tumans = Tuman.objects.all()
    status_choices = ApplicationCreate.STATUS_CHOICES
    hujumturi = HujumTuri.objects.all()

    paginator = Paginator(application, 15)
    page = request.GET.get('page')

    try:
        applications = paginator.page(page)
    except PageNotAnInteger:
        applications = paginator.page(1)
    except EmptyPage:
        applications = paginator.page(paginator.num_pages)


    ctx = {'applications': applications, 'tumans': tumans, 'status_choices': status_choices, 'hujumturi': hujumturi}
    return render(request, 'dashboard/tablitsa.html', ctx)


@login_decorator
def export_to_excel(request):
    # Olingan arizalar
    applications = Application.objects.select_related('hujumturi', 'district', 'app_create').all()

    wb = Workbook()
    ws = wb.active

    # Styles
    header_style = Font(bold=True, color='FFFFFF')
    cell_style = Alignment(horizontal='left', vertical='center')
    header_fill = PatternFill(start_color='0072BA', end_color='0072BA', fill_type='solid')

    # Header row
    header_row = ['F.I.Sh', 'Telefon', 'Tug\'ilgan yili', 'Passport', 'Tuman', 'Manzil', 'jins', 'Hujum turi',
                  'Plastik raqami', 'Plastik raqami \ngumondor', 'Gumondor F.I.Sh', 'Gumondor \n telefon raqami',
                  'Pul yechilgan \nsana', 'Pul yechib \nolingan summa', 'Shaxs ishlatgan\n ilova',
                  'Gumondor\n ishlatgan ilova', 'Shaxs yozgan\n qisqa so\'z', 'Status']
    for col_num, value in enumerate(header_row, 1):
        col_letter = get_column_letter(col_num)
        ws[f'{col_letter}1'] = value
        ws[f'{col_letter}1'].font = header_style
        ws[f'{col_letter}1'].alignment = cell_style
        ws[f'{col_letter}1'].fill = header_fill

    # Data rows
    for row_num, app in enumerate(applications, 2):
        ws[f'A{row_num}'] = f'{app.first_name} {app.last_name} {app.surname}'
        ws[f'B{row_num}'] = app.phone
        ws[f'C{row_num}'] = app.birthday
        ws[f'D{row_num}'] = app.passport_serial
        ws[f'E{row_num}'] = app.district.tuman_name if app.district else ''
        ws[f'F{row_num}'] = app.address
        ws[f'G{row_num}'] = app.gender
        ws[f'H{row_num}'] = app.hujumturi.hujum_name if app.hujumturi else ''
        ws[f'I{row_num}'] = app.plastikraqam_ozi
        ws[f'J{row_num}'] = app.plastikraqam_gumondor
        ws[f'K{row_num}'] = app.full_name_gumondor
        ws[f'L{row_num}'] = app.phone_gumondor
        ws[f'M{row_num}'] = app.vaqt
        ws[f'N{row_num}'] = app.summa
        ws[f'O{row_num}'] = app.ilova
        ws[f'P{row_num}'] = app.ilova_gumondor
        ws[f'Q{row_num}'] = app.text
        ws[f'R{row_num}'] = app.app_create.status if app.app_create else ''

    # Apply borders to all cells
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.border = Border(left=Side(style='thin', color='000000'),
                                 right=Side(style='thin', color='000000'),
                                 top=Side(style='thin', color='000000'),
                                 bottom=Side(style='thin', color='000000'))

    # Excel faylini HTTP javob qilib qaytarish
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=arizalar.xlsx'
    response.write(save_virtual_workbook(wb))

    return response


@login_decorator
def export_to_word(request, pk):
    # Odatiy talablar uchun so'rovni oling
    applications = Application.objects.select_related('hujumturi', 'district', 'app_create').get(pk=pk)

    # Create a new Word document
    document = Document()

    # Add a title to the document
    document.add_heading('Arizalar Ro`yxati', level=1)

    # Add a table to the document
    table = document.add_table(rows=1, cols=21)
    table.autofit = True

    # Add column headers to the table
    column_headers = ['Ism familiya', 'Telefon no`mer', 'Tug`ilgan yili', 'Passpor seriyasi', 'Tuman', 'Manzil',
                      'Jins', 'Hujum turi', 'Plastik raqami', 'Plastik raqami gumondor', 'Gumondor ismi familiyasi',
                      'Gumondor telefon raqami', 'Pul yechilgan sana', 'Yechib olingan summa',
                      'Shaxs ishlatgan ilova', 'Gumodor ishlatgan ilova', 'Shaxs yozgan qisqa mazmuni', 'Status']
    for col_num, header_text in enumerate(column_headers):
        table.cell(0, col_num).text = header_text

    # Add data rows to the table
    for row_num, application in enumerate(applications):
        row = table.add_row()
        row.cells[0].text = f"{application.first_name} {application.last_name} {application.surname}"
        row.cells[1].text = application.phone
        row.cells[2].text = str(application.birthday)
        row.cells[3].text = application.passport_serial
        row.cells[4].text = application.district.tuman_name
        row.cells[5].text = application.address
        row.cells[6].text = application.gender
        row.cells[7].text = application.hujumturi.hujum_name
        row.cells[8].text = application.plastikraqam_ozi
        row.cells[9].text = application.plastikraqam_gumondor
        row.cells[10].text = application.full_name_gumondor
        row.cells[11].text = application.phone_gumondor
        row.cells[12].text = str(application.vaqt)
        row.cells[13].text = f"{application.summa} so'm"
        row.cells[15].text = application.ilova
        row.cells[16].text = application.ilova_gumondor
        row.cells[17].text = application.text
        row.cells[18].text = application.app_create.status
        # Add actions (view, edit, download) links to the last column
        actions = row.cells[20].paragraphs[0]
        actions.add_run(f"{application.id}")  # Assuming you want to link to the application details page

    # Create a response with the Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=shaxs.docx'
    document.save(response)

    return response
